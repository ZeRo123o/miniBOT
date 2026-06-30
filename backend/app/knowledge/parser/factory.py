from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


SUPPORTED_MARKDOWN_EXTENSIONS = {
    ".md",
    ".markdown",
    ".txt",
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
}


def parse_document_to_markdown(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_MARKDOWN_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")

    if suffix in {".md", ".markdown", ".txt"}:
        return _decode_text(content)
    if suffix == ".pdf":
        return _pdf_to_markdown(content)
    if suffix == ".docx":
        return _docx_to_markdown(content)
    if suffix == ".xlsx":
        return _xlsx_to_markdown(content)
    if suffix == ".csv":
        return _csv_to_markdown(content)

    raise ValueError(f"Unsupported document type: {suffix or 'unknown'}")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _pdf_to_markdown(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    parts = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"## Page {index}\n\n{text.strip()}")
    return "\n\n".join(parts).strip()


def _docx_to_markdown(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("heading"):
            level = _heading_level(style_name)
            parts.append(f"{'#' * level} {text}")
        else:
            parts.append(text)

    for table in document.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        table_markdown = _rows_to_markdown(rows)
        if table_markdown:
            parts.append(table_markdown)
    return "\n\n".join(parts).strip()


def _heading_level(style_name: str) -> int:
    digits = "".join(char for char in style_name if char.isdigit())
    if not digits:
        return 2
    return max(1, min(int(digits), 6))


def _xlsx_to_markdown(content: bytes) -> str:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    sheets = []
    for sheet in workbook.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                rows.append(values)
        table_markdown = _rows_to_markdown(rows)
        if table_markdown:
            sheets.append(f"## {sheet.title}\n\n{table_markdown}")
    return "\n\n".join(sheets).strip()


def _csv_to_markdown(content: bytes) -> str:
    import csv

    text = _decode_text(content)
    rows = list(csv.reader(text.splitlines()))
    return _rows_to_markdown(rows)


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    max_columns = max(len(row) for row in rows)
    normalized = [row + [""] * (max_columns - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * max_columns
    body = normalized[1:]
    markdown_rows = [header, separator, *body]
    return "\n".join("| " + " | ".join(_escape_table_cell(cell) for cell in row) + " |" for row in markdown_rows)


def _escape_table_cell(value: str) -> str:
    return str(value).replace("|", "\\|").strip()
